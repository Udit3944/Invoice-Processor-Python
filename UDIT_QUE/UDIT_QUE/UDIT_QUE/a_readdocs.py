#import required libraries
import os
import docx
import openpyxl

#load workbook using openpyxl
word_book = openpyxl.Workbook()

#active worksheet of wordbook
work_sheet = word_book.active

#Write titles into Excel file
work_sheet.append(['Invoice ID', 'Total Products Purchased', 'Subtotal', 'Tax', 'Total'])

#define address of approprite folder directory where all docx save
dire = 'C:/Users/UDIT GAMI/Desktop/Sem-2/UDIT_QUE/UDIT_QUE'

#All define products list for counting
products = ["Parka", "Boots", "Snowshoes", "Climbing Rope", "Oxygen Tank", "Ice Pick", "Crampons"]

#iterate through all .docx file and read data
for file in os.listdir(dire):
    if file.endswith('.docx'):
        file_path = os.path.join(dire, file)
        doc = docx.Document(file_path)

        #write invoice_id to worksheet
        invoice_id = doc.paragraphs[0].text

        #intialize all the required variables
        total_products = 0
        subtotal = 0.0
        tax = 0.0
        total = 0.0

        #iterate through all paragraphs possible line by line
        for para in doc.paragraphs[1:]:

            #split paragraphs into lines
            lines = para.text.split('\n')

            #iterate through line for calculating all data
            for line in lines:

                #if : in text then split into two values
                if ':' in line:

                    #key is name of product and val is quantity in string type
                    key, val = line.split(':', 1)

                    #if it is a product then add it's quantity to total number of products purchased
                    if key in products:

                        #val is string converted to int
                        total_products += int(val)
                    
                    #if it is a subtotal then add it to subtotal
                    elif key == 'SUBTOTAL':

                        #val is string converted to float
                        subtotal = float(val)

                    #if it is a tax then add it to tax
                    elif key == 'TAX':
                        tax = float(val)

                    #if it is a total then add it to total
                    elif key == 'TOTAL':
                        total = float(val)

        #write all the extracted data to work_sheet
        work_sheet.append([invoice_id, total_products, subtotal, tax, total])

#Save the data to A2_Ex.xlsx spread sheet
word_book.save('A2_Ex.xlsx')


print("Invoice Document Datasheet created successfully! and file name is A2_Ex")
